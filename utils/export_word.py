from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

def export_analyse_word(titre, contenu, nom_entreprise="", pays="", exercice=""):
    """
    Génère un document Word professionnel SYSCOHADA.
    """
    doc = Document()

    # --- EN-TÊTE ---
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "SMD Consulting — Superviseur IA Comptable SYSCOHADA"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- TITRE PRINCIPAL ---
    titre_doc = doc.add_heading(titre, level=0)
    titre_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre_doc.runs[0]
    run.font.color.rgb = RGBColor(0x1f, 0x77, 0xb4)

    # --- INFOS ---
    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if nom_entreprise:
        run = info.add_run(f"Entreprise : {nom_entreprise}\n")
        run.bold = True
    if pays:
        run = info.add_run(f"Pays : {pays}\n")
        run.bold = True
    if exercice:
        run = info.add_run(f"Exercice : {exercice}\n")
        run.bold = True

    run = info.add_run(f"Date : {datetime.now().strftime('%d/%m/%Y')}\n")
    run = info.add_run("Généré par SMD Consulting — Superviseur IA Comptable SYSCOHADA")
    run.italic = True

    doc.add_paragraph("")

    # --- CONTENU ---
    lignes = contenu.split("\n")
    for ligne in lignes:
        ligne = ligne.strip()
        if not ligne:
            doc.add_paragraph("")
        elif ligne.startswith("# "):
            doc.add_heading(ligne[2:], level=1)
        elif ligne.startswith("## "):
            doc.add_heading(ligne[3:], level=2)
        elif ligne.startswith("### "):
            doc.add_heading(ligne[4:], level=3)
        elif ligne.startswith("- ") or ligne.startswith("• "):
            doc.add_paragraph(ligne[2:], style="List Bullet")
        elif ligne.startswith("**") and ligne.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(ligne.replace("**", ""))
            run.bold = True
        else:
            doc.add_paragraph(ligne)

    # --- PIED DE PAGE ---
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"© {datetime.now().year} SMD Consulting — Document confidentiel — Normes SYSCOHADA/OHADA"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- SAUVEGARDE ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer