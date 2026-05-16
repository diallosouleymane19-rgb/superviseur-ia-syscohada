"""
smd_engine.py — Moteur de calcul fiscal UEMOA/OHADA
SMD Consulting — Superviseur IA Comptable & Fiscal
Version 2.0

Usage:
    from smd_engine import parse_balance, compute_fiscal, PAYS_UEMOA
"""

import openpyxl
import pandas as pd
from io import BytesIO


# ─────────────────────────────────────────────
#  RÉFÉRENTIEL FISCAL UEMOA (8 pays)
# ─────────────────────────────────────────────
PAYS_UEMOA = {
    "Sénégal": {
        "admin": "DGID",
        "tva_taux": 18.0,
        "is_taux": 30.0,
        "is_imf": 0.5,        # % du CA HT minimum
        "imf_min": 500_000,
        "inps_patron": 15.4,  # IPRES 8.4% + CSS 7%
        "inps_sal": 5.6,
        "echeance_tva": "15 du mois suivant",
        "echeance_is": "30 avril N+1",
        "echeance_inps": "15 du mois suivant",
        "comptes_tva_col": ["44311", "4431"],
        "comptes_tva_ded": ["44541", "44542", "4454", "4455"],
        "comptes_is": ["4411", "44111"],
        "comptes_sal": ["661"],
        "comptes_inps": ["4311", "431"],
        "monnaie": "FCFA (XOF)",
    },
    "Côte d'Ivoire": {
        "admin": "DGI",
        "tva_taux": 18.0,
        "is_taux": 25.0,
        "is_imf": 0.5,
        "imf_min": 3_000_000,
        "inps_patron": 14.77,
        "inps_sal": 3.2,
        "echeance_tva": "15 du mois suivant",
        "echeance_is": "30 avril N+1",
        "echeance_inps": "15 du mois suivant",
        "comptes_tva_col": ["4431"],
        "comptes_tva_ded": ["4454", "4455"],
        "comptes_is": ["441"],
        "comptes_sal": ["661"],
        "comptes_inps": ["431"],
        "monnaie": "FCFA (XOF)",
    },
    "Mali": {
        "admin": "DGI",
        "tva_taux": 18.0,
        "is_taux": 30.0,
        "is_imf": 1.0,
        "imf_min": 300_000,
        "inps_patron": 11.5,
        "inps_sal": 3.5,
        "echeance_tva": "20 du mois suivant",
        "echeance_is": "30 avril N+1",
        "echeance_inps": "20 du mois suivant",
        "comptes_tva_col": ["4431"],
        "comptes_tva_ded": ["4454", "4455"],
        "comptes_is": ["441"],
        "comptes_sal": ["661"],
        "comptes_inps": ["431"],
        "monnaie": "FCFA (XOF)",
    },
    "Burkina Faso": {
        "admin": "DGI",
        "tva_taux": 18.0,
        "is_taux": 27.5,
        "is_imf": 1.0,
        "imf_min": 500_000,
        "inps_patron": 16.0,
        "inps_sal": 5.5,
        "echeance_tva": "20 du mois suivant",
        "echeance_is": "30 avril N+1",
        "echeance_inps": "Trimestrielle",
        "comptes_tva_col": ["4431"],
        "comptes_tva_ded": ["4454", "4455"],
        "comptes_is": ["441"],
        "comptes_sal": ["661"],
        "comptes_inps": ["431"],
        "monnaie": "FCFA (XOF)",
    },
    "Bénin": {
        "admin": "DGI",
        "tva_taux": 18.0,
        "is_taux": 30.0,
        "is_imf": 1.0,
        "imf_min": 300_000,
        "inps_patron": 15.4,
        "inps_sal": 3.6,
        "echeance_tva": "20 du mois suivant",
        "echeance_is": "30 avril N+1",
        "echeance_inps": "Mensuelle/Trimestrielle",
        "comptes_tva_col": ["4431"],
        "comptes_tva_ded": ["4454", "4455"],
        "comptes_is": ["441"],
        "comptes_sal": ["661"],
        "comptes_inps": ["431"],
        "monnaie": "FCFA (XOF)",
    },
    "Togo": {
        "admin": "OTR",
        "tva_taux": 18.0,
        "is_taux": 27.0,
        "is_imf": 1.5,
        "imf_min": 800_000,
        "inps_patron": 17.5,
        "inps_sal": 4.0,
        "echeance_tva": "20 du mois suivant",
        "echeance_is": "30 avril N+1",
        "echeance_inps": "20 du mois suivant",
        "comptes_tva_col": ["4431"],
        "comptes_tva_ded": ["4454", "4455"],
        "comptes_is": ["441"],
        "comptes_sal": ["661"],
        "comptes_inps": ["431"],
        "monnaie": "FCFA (XOF)",
    },
    "Niger": {
        "admin": "DGI",
        "tva_taux": 19.0,
        "is_taux": 30.0,
        "is_imf": 1.0,
        "imf_min": 300_000,
        "inps_patron": 15.4,
        "inps_sal": 3.0,
        "echeance_tva": "15 du mois suivant",
        "echeance_is": "30 avril N+1",
        "echeance_inps": "Trimestrielle",
        "comptes_tva_col": ["4431"],
        "comptes_tva_ded": ["4454", "4455"],
        "comptes_is": ["441"],
        "comptes_sal": ["661"],
        "comptes_inps": ["431"],
        "monnaie": "FCFA (XOF)",
    },
    "Guinée-Bissau": {
        "admin": "DGCI",
        "tva_taux": 19.0,   # IGV jusqu'au 31/12/2024, TVA depuis 01/01/2025
        "is_taux": 25.0,    # IRPC
        "is_imf": 0.0,
        "imf_min": 0,
        "inps_patron": 14.0,
        "inps_sal": 8.0,
        "echeance_tva": "20 du mois suivant",
        "echeance_is": "31 mars N+1",
        "echeance_inps": "20 du mois suivant",
        "comptes_tva_col": ["443100", "4431"],
        "comptes_tva_ded": ["443201", "443600", "444100", "445201", "445202", "445203", "4454", "4455"],
        "comptes_is": ["441000", "441"],
        "comptes_sal": ["661100", "661200", "661800", "661"],
        "comptes_inps": ["431100", "431"],
        "monnaie": "FCFA (XOF)",
    },
}


# ─────────────────────────────────────────────
#  LECTURE DE LA BALANCE SYSCOHADA (6 colonnes)
# ─────────────────────────────────────────────
def parse_balance(file_obj) -> pd.DataFrame:
    """
    Lit une balance des comptes à 6 colonnes (format Sage/SYSCOHADA).
    Détecte automatiquement les colonnes Débit/Crédit et les soldes.

    Retourne un DataFrame avec colonnes :
        compte, libelle, mvt_deb, mvt_cre, sold_deb, sold_cre
    """
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    rows = [list(r) for r in ws.iter_rows(values_only=True)]

    records = []
    for r in rows:
        c0 = r[0]
        if c0 is None:
            continue
        c0_str = str(c0).strip().replace('.0', '')
        # Accepter les comptes numériques de 3+ chiffres
        if not c0_str.replace('.', '').isdigit() or len(c0_str.replace('.', '')) < 3:
            continue

        compte = c0_str if not c0_str.endswith('.') else c0_str[:-1]
        libelle = str(r[2]) if r[2] else (str(r[1]) if r[1] else "")

        def safe_float(v):
            try:
                return float(v) if v is not None else 0.0
            except (ValueError, TypeError):
                return 0.0

        # Colonnes standard balance 6 colonnes (indices empiriques Sage)
        mvt_deb  = safe_float(r[16]) if len(r) > 16 else 0.0
        mvt_cre  = safe_float(r[19]) if len(r) > 19 else 0.0
        sold_deb = safe_float(r[22]) if len(r) > 22 else 0.0
        sold_cre = safe_float(r[25]) if len(r) > 25 else 0.0

        # Fallback si les colonnes standards sont vides — essayer d'autres positions
        if mvt_deb == 0 and mvt_cre == 0:
            # Chercher les colonnes avec des valeurs numériques
            nums = [(i, safe_float(v)) for i, v in enumerate(r) if i > 2 and safe_float(v) != 0]
            if len(nums) >= 4:
                mvt_deb  = nums[0][1]
                mvt_cre  = nums[1][1]
                sold_deb = nums[2][1]
                sold_cre = nums[3][1]

        records.append({
            "compte":   compte,
            "libelle":  libelle.strip(),
            "mvt_deb":  mvt_deb,
            "mvt_cre":  mvt_cre,
            "sold_deb": sold_deb,
            "sold_cre": sold_cre,
        })

    return pd.DataFrame(records)


def _sum_comptes(df: pd.DataFrame, prefixes: list, col: str) -> float:
    """Somme une colonne pour tous les comptes commençant par les préfixes donnés."""
    total = 0.0
    for prefix in prefixes:
        mask = df["compte"].str.startswith(prefix)
        total += df.loc[mask, col].sum()
    return total


# ─────────────────────────────────────────────
#  MOTEUR DE CALCUL FISCAL
# ─────────────────────────────────────────────
def compute_fiscal(df: pd.DataFrame, pays: str, entreprise: str = "", periode: str = "") -> dict:
    """
    Calcule tous les indicateurs fiscaux à partir du DataFrame de balance.

    Args:
        df:          DataFrame retourné par parse_balance()
        pays:        Nom du pays (clé de PAYS_UEMOA)
        entreprise:  Nom de l'entreprise (optionnel)
        periode:     Période (ex: "Exercice 2019")

    Returns:
        dict avec toutes les métriques, alertes et recommandations
    """
    p = PAYS_UEMOA.get(pays, PAYS_UEMOA["Guinée-Bissau"])

    def sc(prefixes): return _sum_comptes(df, prefixes, "sold_cre")
    def sd(prefixes): return _sum_comptes(df, prefixes, "sold_deb")
    def md(prefixes): return _sum_comptes(df, prefixes, "mvt_deb")
    def mc(prefixes): return _sum_comptes(df, prefixes, "mvt_cre")

    # ── Chiffre d'affaires ──
    ca = sc(["701", "702", "703", "706", "707", "70"])

    # ── Charges ──
    charges_total = sd(["6"])

    # ── Résultat ──
    resultat = ca - charges_total

    # ── TVA / IGV ──
    tva_collectee  = sc(p["comptes_tva_col"])
    tva_deductible = sd(p["comptes_tva_ded"])
    tva_nette      = tva_collectee - tva_deductible

    # Détail TVA déductible par compte
    tva_detail = []
    for compte_prefix in p["comptes_tva_ded"]:
        mask = df["compte"].str.startswith(compte_prefix)
        sub = df[mask]
        for _, row in sub.iterrows():
            if row["sold_deb"] > 0:
                tva_detail.append({
                    "compte":  row["compte"],
                    "libelle": row["libelle"],
                    "montant": row["sold_deb"],
                })

    # ── IS / IRPC ──
    is_theorique = max(resultat * p["is_taux"] / 100, 0)
    is_paye      = md(p["comptes_is"])

    # Réintégrations courantes (dons, pénalités)
    reintegrations = sd(["658"]) + sd(["647"])  # dons + pénalités
    resultat_fiscal = resultat + reintegrations
    is_corrige      = max(resultat_fiscal * p["is_taux"] / 100, 0)
    is_solde        = is_corrige - is_paye

    # IMF (minimum forfaitaire)
    if p["is_imf"] > 0:
        imf = max(ca * p["is_imf"] / 100, p["imf_min"])
        is_a_payer = max(is_corrige, imf)
    else:
        imf = 0
        is_a_payer = is_corrige

    # ── INPS / CNSS ──
    sal_comptes = p["comptes_sal"]
    salaires_bruts = sd(sal_comptes)

    inps_patron_du  = salaires_bruts * p["inps_patron"] / 100
    inps_sal_du     = salaires_bruts * p["inps_sal"] / 100
    inps_total_du   = inps_patron_du + inps_sal_du
    inps_paye       = md(p["comptes_inps"])
    inps_provisionne = mc(p["comptes_inps"])
    inps_solde_cpte = sd(p["comptes_inps"])  # débiteur = trop-payé
    inps_ecart      = inps_total_du - inps_paye

    # ── Autres risques ──
    aci              = sd(["441200", "4412"])
    cca              = sc(["47119", "4711"])
    stock_negatif    = sc(["311", "31"])  # stock créditeur = anomalie
    avances_personnel = sd(["4211", "421100"])
    creances_dout    = sd(["41110", "411106"])
    dons             = sd(["658"])
    penalites        = sd(["647"])
    decouvert_banque = sc(["521", "52"])  # banque créditeur = découvert
    capital          = sc(["101"])
    deficit_reporte  = sd(["121"])

    # ── Ratios ──
    ratio_tva_ca     = (tva_nette / ca * 100) if ca > 0 else 0
    pression_fiscale = ((is_a_payer + tva_nette + inps_total_du) / ca * 100) if ca > 0 else 0
    ratio_capital_ca = (capital / ca * 100) if ca > 0 else 0

    # ── Alertes ──
    alertes = []

    if tva_nette > 0:
        niveau = "ÉLEVÉ" if tva_nette > ca * 0.05 else "MODÉRÉ"
        alertes.append({
            "titre": f"TVA/IGV nette non reversée",
            "montant": tva_nette,
            "niveau": niveau,
            "detail": f"TVA collectée {_fmt(tva_collectee)} − Déductible {_fmt(tva_deductible)} = {_fmt(tva_nette)} à reverser à la {p['admin']}",
            "action": f"Déclarer et payer avant le {p['echeance_tva']}. Rapprocher avec les déclarations mensuelles déposées.",
        })

    if is_solde > 0:
        niveau = "MODÉRÉ" if is_solde < ca * 0.02 else "ÉLEVÉ"
        alertes.append({
            "titre": "IS/IRPC résiduel après réintégrations",
            "montant": is_solde,
            "niveau": niveau,
            "detail": f"Résultat fiscal corrigé {_fmt(resultat_fiscal)} × {p['is_taux']}% = {_fmt(is_corrige)} − Acomptes {_fmt(is_paye)} = {_fmt(is_solde)}",
            "action": f"Régulariser avant le {p['echeance_is']}. Vérifier les réintégrations fiscales.",
        })

    if inps_ecart > inps_total_du * 0.05:
        alertes.append({
            "titre": "Cotisations sociales — écart à vérifier",
            "montant": inps_ecart,
            "niveau": "MODÉRÉ",
            "detail": f"INPS/CNSS théorique (22%) : {_fmt(inps_total_du)} — Versé : {_fmt(inps_paye)} — Écart : {_fmt(inps_ecart)}",
            "action": f"Vérifier les bulletins de paie et les reversements {p['echeance_inps']}.",
        })

    if aci > 0:
        alertes.append({
            "titre": "ACI / Avances douanières non apurées",
            "montant": aci,
            "niveau": "MODÉRÉ",
            "detail": f"Avances versées aux services douaniers non encore justifiées : {_fmt(aci)}",
            "action": "Constituer le dossier douanier complet (DAU, factures fournisseurs, quittances).",
        })

    if cca > 0:
        niveau = "ÉLEVÉ" if cca > ca * 0.05 else "MODÉRÉ"
        alertes.append({
            "titre": "Compte courant actionnaires non formalisé",
            "montant": cca,
            "niveau": niveau,
            "detail": f"Avances actionnaires sans convention formalisée : {_fmt(cca)}",
            "action": "Formaliser par une convention écrite avec taux d'intérêt de marché.",
        })

    if stock_negatif > 0:
        alertes.append({
            "titre": "Stock comptablement négatif — anomalie",
            "montant": stock_negatif,
            "niveau": "ÉLEVÉ",
            "detail": f"Solde créditeur sur compte de stock : {_fmt(stock_negatif)} — anomalie comptable.",
            "action": "Procéder à un inventaire physique. Corriger les imputations.",
        })

    if dons > 0:
        alertes.append({
            "titre": "Dons — charges non déductibles à réintégrer",
            "montant": dons,
            "niveau": "MODÉRÉ",
            "detail": f"Dons comptabilisés en charges : {_fmt(dons)} — généralement non déductibles fiscalement.",
            "action": "Réintégrer dans la base IS/IRPC sauf justification par texte fiscal local.",
        })

    # Exposition totale
    exposition_totale = sum(a["montant"] for a in alertes)

    return {
        # ── Identité ──
        "entreprise":      entreprise,
        "pays":            pays,
        "periode":         periode,
        "admin":           p["admin"],
        "monnaie":         p["monnaie"],

        # ── CA & Résultat ──
        "ca":              ca,
        "charges":         charges_total,
        "resultat":        resultat,
        "reintegrations":  reintegrations,
        "resultat_fiscal": resultat_fiscal,

        # ── TVA / IGV ──
        "tva_taux":        p["tva_taux"],
        "tva_collectee":   tva_collectee,
        "tva_deductible":  tva_deductible,
        "tva_nette":       tva_nette,
        "tva_detail":      tva_detail,

        # ── IS / IRPC ──
        "is_taux":         p["is_taux"],
        "is_theorique":    is_theorique,
        "is_corrige":      is_corrige,
        "is_paye":         is_paye,
        "is_solde":        is_solde,
        "imf":             imf,
        "is_a_payer":      is_a_payer,

        # ── INPS / CNSS ──
        "inps_patron_taux":  p["inps_patron"],
        "inps_sal_taux":     p["inps_sal"],
        "salaires_bruts":    salaires_bruts,
        "inps_patron_du":    inps_patron_du,
        "inps_sal_du":       inps_sal_du,
        "inps_total_du":     inps_total_du,
        "inps_paye":         inps_paye,
        "inps_provisionne":  inps_provisionne,
        "inps_solde":        inps_solde_cpte,
        "inps_ecart":        inps_ecart,

        # ── Autres ──
        "aci":               aci,
        "cca":               cca,
        "stock_negatif":     stock_negatif,
        "avances_personnel": avances_personnel,
        "creances_dout":     creances_dout,
        "dons":              dons,
        "penalites":         penalites,
        "decouvert_banque":  decouvert_banque,
        "capital":           capital,
        "deficit_reporte":   deficit_reporte,

        # ── Ratios ──
        "ratio_tva_ca":      ratio_tva_ca,
        "pression_fiscale":  pression_fiscale,
        "ratio_capital_ca":  ratio_capital_ca,

        # ── Alertes & Exposition ──
        "alertes":           alertes,
        "exposition_totale": exposition_totale,

        # ── Échéances ──
        "echeance_tva":    p["echeance_tva"],
        "echeance_is":     p["echeance_is"],
        "echeance_inps":   p["echeance_inps"],
    }


def _fmt(n: float) -> str:
    """Formate un nombre en FCFA avec séparateurs."""
    return f"{int(abs(n)):,}".replace(",", " ") + " FCFA"


# ─────────────────────────────────────────────
#  GÉNÉRATEUR EXCEL (tableau de bord)
# ─────────────────────────────────────────────
def generate_excel(result: dict) -> BytesIO:
    """Génère le tableau de bord fiscal en Excel et retourne un BytesIO."""
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tableau de Bord Fiscal"

    # Couleurs
    BLUE_D, BLUE_L = "1A5276", "D4E6F1"
    RED_D,  RED_L  = "C0392B", "FADBD8"
    ORG_D,  ORG_L  = "D35400", "FDEBD0"
    GRN_D,  GRN_L  = "1E8449", "D5F5E3"
    GREY, WHITE    = "F2F3F4", "FFFFFF"

    def fill(c): return PatternFill("solid", fgColor=c)
    def fnt(bold=False, color="111111", sz=10): return Font(bold=bold, color=color, size=sz, name="Arial")
    def aln(h="left", v="center"): return Alignment(horizontal=h, vertical=v, wrap_text=True)
    thin = Side(style='thin', color='BDBDBD')
    def brd(): return Border(left=thin, right=thin, top=thin, bottom=thin)

    def write(row, col, val, bg=WHITE, fg="111111", bold=False, sz=10, h="left", num_fmt=None):
        c = ws.cell(row=row, column=col, value=val)
        c.fill = fill(bg); c.font = fnt(bold, fg, sz)
        c.alignment = aln(h); c.border = brd()
        if num_fmt: c.number_format = num_fmt
        return c

    # Largeurs colonnes
    for i, w in enumerate([3, 38, 22, 18, 3], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    r = 1
    # Titre
    ws.merge_cells(f'B{r}:D{r}')
    c = ws.cell(row=r, column=2, value="TABLEAU DE BORD FISCAL MENSUEL")
    c.fill = fill(BLUE_D); c.font = Font(bold=True, color="FFFFFF", size=16, name="Arial")
    c.alignment = aln("center"); ws.row_dimensions[r].height = 32; r += 1

    ws.merge_cells(f'B{r}:D{r}')
    c = ws.cell(row=r, column=2, value=f"{result['entreprise']}  |  {result['pays']}  |  {result['periode']}")
    c.fill = fill(BLUE_L); c.font = Font(bold=True, color=BLUE_D, size=11, name="Arial")
    c.alignment = aln("center"); ws.row_dimensions[r].height = 22; r += 1; r += 1

    # KPIs
    kpis = [
        ("Chiffre d'Affaires", result["ca"], BLUE_D, BLUE_L),
        ("Résultat Net", result["resultat"], GRN_D, GRN_L),
        ("TVA/IGV Nette Due", result["tva_nette"], RED_D, RED_L),
        ("IS/IRPC Résiduel", result["is_solde"], ORG_D, ORG_L),
    ]
    kpi_cols = [2, 3]  # B et C — 2 colonnes de KPIs
    for i, (label, val, fd, fl) in enumerate(kpis):
        row_label = r + (i // 2) * 2
        col = 2 + (i % 2)
        c1 = ws.cell(row=row_label, column=col, value=label)
        c1.fill = fill(fl); c1.font = Font(bold=True, color=fd, size=10, name="Arial")
        c1.alignment = aln("center"); c1.border = brd()
        ws.row_dimensions[row_label].height = 18
        c2 = ws.cell(row=row_label+1, column=col, value=int(val))
        c2.fill = fill(fl); c2.font = Font(bold=True, color=fd, size=12, name="Arial")
        c2.alignment = aln("center"); c2.border = brd()
        c2.number_format = '#,##0'
        ws.row_dimensions[row_label+1].height = 22

    r += 6

    def section(title, bg_color, rows_data, headers):
        nonlocal r
        ws.merge_cells(f'B{r}:D{r}')
        c = ws.cell(row=r, column=2, value=f"  {title}")
        c.fill = fill(bg_color); c.font = Font(bold=True, color="FFFFFF", size=11, name="Arial")
        c.alignment = aln("left"); ws.row_dimensions[r].height = 22; r += 1
        for hi, h in enumerate(headers):
            write(r, hi+2, h, bg=bg_color, fg="FFFFFF", bold=True, h="center")
        ws.row_dimensions[r].height = 18; r += 1
        for i, row_d in enumerate(rows_data):
            bg = WHITE if i % 2 == 0 else GREY
            for ci, v in enumerate(row_d):
                h = "right" if isinstance(v, (int, float)) else "left"
                c = write(r, ci+2, int(v) if isinstance(v, float) and v == int(v) else v, bg=bg, h=h)
                if isinstance(v, (int, float)):
                    c.number_format = '#,##0'
            ws.row_dimensions[r].height = 18; r += 1
        r += 1  # espace

    # Section TVA
    section("TVA / IGV", RED_D, [
        ("TVA collectée",       result["tva_collectee"],  "Passif fiscal"),
        ("TVA déductible total", -result["tva_deductible"], "Crédit"),
        ("TVA NETTE DUE",       result["tva_nette"],      "🔴 À REVERSER"),
    ], ["Élément", "Montant (FCFA)", "Statut"])

    # Section IS
    section("IS / IRPC", ORG_D, [
        ("Résultat comptable", result["resultat"], "Base départ"),
        ("+ Réintégrations",  result["reintegrations"], "Dons + pénalités"),
        ("= Résultat fiscal",  result["resultat_fiscal"], "Base IS"),
        (f"IS ({result['is_taux']}%)",  result["is_corrige"], "Impôt calculé"),
        ("IS déjà versé",     -result["is_paye"], "Acomptes"),
        ("IS RÉSIDUEL",       result["is_solde"], "🟠 À RÉGULARISER"),
    ], ["Calcul", "Montant (FCFA)", "Note"])

    # Section INPS
    section("INPS / CNSS", BLUE_D, [
        ("Masse salariale brute",  result["salaires_bruts"], "Base"),
        (f"Part patronale ({result['inps_patron_taux']}%)", result["inps_patron_du"], "Employeur"),
        (f"Part salariale ({result['inps_sal_taux']}%)",  result["inps_sal_du"], "Salarié"),
        ("TOTAL INPS attendu",    result["inps_total_du"], "À verser"),
        ("INPS versé",            -result["inps_paye"], "Réel"),
        ("Écart à vérifier",      result["inps_ecart"], "🔵 Vérifier"),
    ], ["Élément", "Montant (FCFA)", "Statut"])

    # Section Alertes
    if result["alertes"]:
        section("ALERTES FISCALES", "6C3483", [
            (a["titre"], int(a["montant"]), a["niveau"])
            for a in result["alertes"]
        ], ["Risque", "Montant (FCFA)", "Niveau"])

    # Total exposition
    ws.merge_cells(f'B{r}:D{r}')
    c = ws.cell(row=r, column=2, value=f"EXPOSITION TOTALE ESTIMÉE : {_fmt(result['exposition_totale'])}")
    c.fill = fill(RED_D); c.font = Font(bold=True, color="FFFFFF", size=13, name="Arial")
    c.alignment = aln("center"); ws.row_dimensions[r].height = 28

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


# ─────────────────────────────────────────────
#  GÉNÉRATEUR WORD (rapport risque fiscal)
# ─────────────────────────────────────────────
def generate_word(result: dict) -> BytesIO:
    """Génère le rapport d'analyse du risque fiscal en .docx et retourne un BytesIO."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        # Si python-docx n'est pas installé, retourner None
        return None

    doc = DocxDocument()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def rgb(hex_color):
        h = hex_color.lstrip('#')
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def add_heading(text, level=1, color="1A5276"):
        h = doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = rgb(color)
        h.runs[0].font.name = "Arial"
        return h

    def add_para(text, bold=False, color="222222", size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.color.rgb = rgb(color)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        return p

    def add_table_row(table, cells, bgs=None, bold_flags=None):
        row = table.add_row()
        for i, (cell_val, cell) in enumerate(zip(cells, row.cells)):
            cell.text = str(cell_val) if not isinstance(cell_val, float) else _fmt(cell_val)
            run = cell.paragraphs[0].runs
            if run:
                run[0].font.name = "Arial"
                run[0].font.size = Pt(9)
                if bold_flags and bold_flags[i]:
                    run[0].bold = True
            if bgs and i < len(bgs) and bgs[i]:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), bgs[i])
                tcPr.append(shd)

    # Titre
    doc.add_heading("", 0)  # clear default
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("RAPPORT D'ANALYSE DU RISQUE FISCAL")
    run.bold = True; run.font.size = Pt(20); run.font.color.rgb = rgb("1A5276"); run.font.name = "Arial"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"{result['entreprise']}  —  {result['pays']}  —  {result['periode']}")
    r2.font.size = Pt(13); r2.font.color.rgb = rgb("444444"); r2.font.name = "Arial"

    doc.add_paragraph("")

    # Résumé
    add_heading("1. Résumé Exécutif")
    add_para(
        f"Entreprise : {result['entreprise']}  |  Pays : {result['pays']}  |  "
        f"Administration : {result['admin']}  |  Période : {result['periode']}"
    )
    add_para(
        f"Chiffre d'affaires : {_fmt(result['ca'])}  |  "
        f"Résultat net : {_fmt(result['resultat'])}  |  "
        f"Exposition fiscale totale estimée : {_fmt(result['exposition_totale'])}",
        bold=True, color="C0392B"
    )
    doc.add_paragraph("")

    # Alertes
    add_heading("2. Alertes Fiscales Identifiées")
    if result["alertes"]:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        for i, txt in enumerate(["Risque", "Montant", "Niveau"]):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].runs[0].font.name = "Arial"

        for a in result["alertes"]:
            row = tbl.add_row().cells
            row[0].text = a["titre"]
            row[1].text = _fmt(a["montant"])
            row[2].text = a["niveau"]
            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].runs[0].font.name = "Arial"

    doc.add_paragraph("")

    # TVA
    add_heading("3. Position TVA / IGV", level=2)
    add_para(f"Taux applicable : {result['tva_taux']} %  —  Administration : {result['admin']}")
    add_para(f"TVA collectée :   {_fmt(result['tva_collectee'])}", bold=True)
    add_para(f"TVA déductible :  {_fmt(result['tva_deductible'])}")
    add_para(f"TVA NETTE DUE :   {_fmt(result['tva_nette'])}", bold=True, color="C0392B")
    add_para(f"Échéance déclaration : {result['echeance_tva']}")
    doc.add_paragraph("")

    # IS
    add_heading("4. IS / IRPC", level=2)
    add_para(f"Taux IS : {result['is_taux']} %")
    add_para(f"Résultat comptable :  {_fmt(result['resultat'])}")
    add_para(f"Réintégrations :      {_fmt(result['reintegrations'])}")
    add_para(f"Résultat fiscal :     {_fmt(result['resultat_fiscal'])}", bold=True)
    add_para(f"IS théorique :        {_fmt(result['is_corrige'])}")
    add_para(f"IS déjà versé :       {_fmt(result['is_paye'])}")
    add_para(f"IS RÉSIDUEL :         {_fmt(result['is_solde'])}", bold=True, color="D35400")
    add_para(f"Échéance déclaration : {result['echeance_is']}")
    doc.add_paragraph("")

    # INPS
    add_heading("5. INPS / CNSS", level=2)
    add_para(f"Part patronale : {result['inps_patron_taux']} %  |  Part salariale : {result['inps_sal_taux']} %")
    add_para(f"Salaires bruts :       {_fmt(result['salaires_bruts'])}")
    add_para(f"INPS total attendu :   {_fmt(result['inps_total_du'])}", bold=True)
    add_para(f"INPS versé :           {_fmt(result['inps_paye'])}")
    add_para(f"Écart à vérifier :     {_fmt(result['inps_ecart'])}", bold=True, color="1A5276")
    add_para(f"Échéance : {result['echeance_inps']}")
    doc.add_paragraph("")

    # Exposition totale
    add_heading("6. Exposition Fiscale Totale")
    add_para(
        f"EXPOSITION TOTALE ESTIMÉE : {_fmt(result['exposition_totale'])}",
        bold=True, color="C0392B", size=14
    )
    add_para("(hors pénalités et majorations de retard)", color="666666")
    doc.add_paragraph("")

    # Avertissement
    add_para(
        "Avertissement : Ce rapport est produit sur la base de la balance des comptes fournie. "
        "Il ne constitue pas un audit légal. SMD Consulting recommande de consulter un expert-comptable "
        "ou avocat fiscaliste local avant toute démarche auprès de l'administration fiscale.",
        color="888888"
    )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
