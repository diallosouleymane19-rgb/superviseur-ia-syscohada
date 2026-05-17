"""
smd_streamlit.py — Pages Streamlit SMD Consulting
Tableau de Bord Fiscal + Analyse du Risque Fiscal
Version 2.0 — UEMOA / OHADA

Intégration dans votre app existante :
---------------------------------------
    # Dans votre app.py principal :
    from smd_streamlit import page_dashboard, page_risque_fiscal

    # Exemple avec st.sidebar pour la navigation :
    page = st.sidebar.radio("Module", ["Tableau de Bord Fiscal", "Analyse Risque Fiscal"])
    if page == "Tableau de Bord Fiscal":
        page_dashboard()
    elif page == "Analyse Risque Fiscal":
        page_risque_fiscal()

Dépendances :
    pip install streamlit openpyxl pandas plotly python-docx
"""

import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from io import BytesIO

# Import du moteur de calcul
from smd_engine import (
    parse_balance, compute_fiscal, generate_excel, generate_word,
    PAYS_UEMOA, _fmt
)

# ─────────────────────────────────────────────
#  CONSTANTES UI
# ─────────────────────────────────────────────
COULEUR_BLEU   = "#1A5276"
COULEUR_ROUGE  = "#C0392B"
COULEUR_ORANGE = "#D35400"
COULEUR_VERT   = "#1E8449"
COULEUR_GRIS   = "#F2F3F4"

NIVEAU_COULEUR = {
    "ÉLEVÉ":   COULEUR_ROUGE,
    "MODÉRÉ":  COULEUR_ORANGE,
    "FAIBLE":  COULEUR_VERT,
    "INFO":    COULEUR_BLEU,
}

NIVEAU_EMOJI = {
    "ÉLEVÉ":  "🔴",
    "MODÉRÉ": "🟠",
    "FAIBLE": "🟢",
    "INFO":   "🔵",
}

# CSS global injecté une seule fois
_CSS = """
<style>
/* Bandeau titre SMD */
.smd-header {
    background: linear-gradient(135deg, #1A5276 0%, #2E86C1 100%);
    color: white;
    padding: 1.2rem 1.8rem;
    border-radius: 10px;
    margin-bottom: 1.5rem;
}
.smd-header h1 { color: white !important; margin: 0; font-size: 1.6rem; }
.smd-header p  { margin: 0.2rem 0 0; opacity: 0.85; font-size: 0.95rem; }

/* Cartes KPI */
.kpi-card {
    background: white;
    border-radius: 8px;
    padding: 1rem;
    border-left: 5px solid #1A5276;
    box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    margin-bottom: 0.5rem;
}
.kpi-card.rouge  { border-left-color: #C0392B; }
.kpi-card.orange { border-left-color: #D35400; }
.kpi-card.vert   { border-left-color: #1E8449; }
.kpi-label { font-size: 0.78rem; color: #666; font-weight: 600; text-transform: uppercase; }
.kpi-value { font-size: 1.3rem; font-weight: 700; color: #222; margin-top: 0.2rem; }
.kpi-value.rouge  { color: #C0392B; }
.kpi-value.orange { color: #D35400; }
.kpi-value.vert   { color: #1E8449; }

/* Badges niveau */
.badge {
    display: inline-block;
    padding: 0.2rem 0.7rem;
    border-radius: 20px;
    font-size: 0.78rem;
    font-weight: 700;
    text-transform: uppercase;
}
.badge.eleve  { background: #FADBD8; color: #C0392B; }
.badge.modere { background: #FDEBD0; color: #D35400; }
.badge.faible { background: #D5F5E3; color: #1E8449; }
.badge.info   { background: #D4E6F1; color: #1A5276; }

/* Boîtes d'alerte */
.alerte-box {
    border-radius: 8px;
    padding: 0.8rem 1.1rem;
    margin-bottom: 0.6rem;
    border-left: 5px solid;
}
.alerte-box.eleve  { background: #FADBD8; border-color: #C0392B; }
.alerte-box.modere { background: #FDEBD0; border-color: #D35400; }
.alerte-box.faible { background: #D5F5E3; border-color: #1E8449; }

/* Exposition totale */
.exposition-box {
    background: #FADBD8;
    border: 2px solid #C0392B;
    border-radius: 10px;
    padding: 1.2rem;
    text-align: center;
    margin: 1rem 0;
}
.exposition-box .label { font-size: 0.9rem; color: #C0392B; font-weight: 600; }
.exposition-box .montant { font-size: 2rem; font-weight: 800; color: #C0392B; }

/* Upload zone */
.upload-zone {
    border: 2px dashed #2E86C1;
    border-radius: 10px;
    padding: 1.5rem;
    text-align: center;
    background: #EBF5FB;
    margin: 1rem 0;
}
</style>
"""


def _inject_css():
    st.markdown(_CSS, unsafe_allow_html=True)


def _kpi(label: str, value: str, couleur: str = "bleu"):
    st.markdown(f"""
    <div class="kpi-card {couleur}">
        <div class="kpi-label">{label}</div>
        <div class="kpi-value {couleur}">{value}</div>
    </div>
    """, unsafe_allow_html=True)


def _badge(niveau: str) -> str:
    cls = {"ÉLEVÉ": "eleve", "MODÉRÉ": "modere", "FAIBLE": "faible"}.get(niveau, "info")
    emoji = NIVEAU_EMOJI.get(niveau, "🔵")
    return f'<span class="badge {cls}">{emoji} {niveau}</span>'


def _alerte_box(titre: str, montant: float, detail: str, action: str, niveau: str):
    cls = {"ÉLEVÉ": "eleve", "MODÉRÉ": "modere", "FAIBLE": "faible"}.get(niveau, "info")
    emoji = NIVEAU_EMOJI.get(niveau, "🔵")
    st.markdown(f"""
    <div class="alerte-box {cls}">
        <strong>{emoji} {titre}</strong> &nbsp; {_badge(niveau)}<br>
        <span style="font-size:0.9rem; color:#333;">{_fmt(montant)}</span><br>
        <small style="color:#555;">{detail}</small><br>
        <small><strong>Action :</strong> {action}</small>
    </div>
    """, unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  GRAPHIQUES PLOTLY
# ─────────────────────────────────────────────
def _chart_tva(result: dict):
    labels = ["TVA collectée", "TVA déductible", "TVA nette due"]
    values = [result["tva_collectee"], result["tva_deductible"], result["tva_nette"]]
    colors = ["#1A5276", "#1E8449", "#C0392B"]
    fig = go.Figure(go.Bar(
        x=labels, y=values, marker_color=colors,
        text=[_fmt(v) for v in values], textposition="outside",
        textfont=dict(size=10)
    ))
    fig.update_layout(
        title=dict(text="Position TVA / IGV", font=dict(size=14, color=COULEUR_BLEU)),
        yaxis_title="FCFA", height=300,
        plot_bgcolor="white", paper_bgcolor="white",
        margin=dict(t=40, b=20, l=20, r=20),
        showlegend=False,
    )
    fig.update_yaxes(gridcolor="#F0F0F0")
    return fig


def _chart_is(result: dict):
    labels = ["Résultat comptable", "Réintégrations", "IS calculé", "IS versé", "IS résiduel"]
    values = [result["resultat"], result["reintegrations"],
              result["is_corrige"], result["is_paye"], result["is_solde"]]
    colors = [COULEUR_BLEU, COULEUR_ORANGE, COULEUR_ROUGE, COULEUR_VERT, COULEUR_ROUGE]
    fig = go.Figure(go.Bar(
        x=labels, y=values, marker_color=colors,
        text=[_fmt(v) for v in values], textposition="outside",
        textfont=dict(size=9)
    ))
    fig.update_layout(
        title=dict(text="IS / IRPC — Calcul fiscal", font=dict(size=14, color=COULEUR_BLEU)),
        yaxis_title="FCFA", height=300,
        plot_bgcolor="white", paper_bgcolor="white",
        margin=dict(t=40, b=20, l=20, r=20),
    )
    fig.update_yaxes(gridcolor="#F0F0F0")
    return fig


def _chart_inps(result: dict):
    fig = go.Figure(go.Bar(
        x=["INPS patronal attendu", "INPS salarial attendu", "Total attendu", "INPS versé"],
        y=[result["inps_patron_du"], result["inps_sal_du"],
           result["inps_total_du"], result["inps_paye"]],
        marker_color=[COULEUR_BLEU, "#2E86C1", COULEUR_ORANGE, COULEUR_VERT],
        text=[_fmt(v) for v in [result["inps_patron_du"], result["inps_sal_du"],
                                 result["inps_total_du"], result["inps_paye"]]],
        textposition="outside", textfont=dict(size=9)
    ))
    fig.update_layout(
        title=dict(text="INPS / CNSS — Cotisations sociales", font=dict(size=14, color=COULEUR_BLEU)),
        yaxis_title="FCFA", height=300,
        plot_bgcolor="white", paper_bgcolor="white",
        margin=dict(t=40, b=20, l=20, r=20),
    )
    fig.update_yaxes(gridcolor="#F0F0F0")
    return fig


def _chart_exposition(result: dict):
    alertes = result["alertes"]
    if not alertes:
        return None
    labels  = [a["titre"][:35] + "…" if len(a["titre"]) > 35 else a["titre"] for a in alertes]
    values  = [a["montant"] for a in alertes]
    colors  = [COULEUR_ROUGE if a["niveau"] == "ÉLEVÉ" else
               COULEUR_ORANGE if a["niveau"] == "MODÉRÉ" else COULEUR_VERT
               for a in alertes]
    fig = go.Figure(go.Bar(
        y=labels, x=values, orientation="h",
        marker_color=colors,
        text=[_fmt(v) for v in values], textposition="outside",
        textfont=dict(size=9)
    ))
    fig.update_layout(
        title=dict(text="Exposition Fiscale par Risque", font=dict(size=14, color=COULEUR_ROUGE)),
        xaxis_title="FCFA", height=max(250, len(alertes) * 45),
        plot_bgcolor="white", paper_bgcolor="white",
        margin=dict(t=40, b=20, l=200, r=80),
        yaxis=dict(autorange="reversed"),
    )
    fig.update_xaxes(gridcolor="#F0F0F0")
    return fig


# ─────────────────────────────────────────────
#  FORMULAIRE D'UPLOAD COMMUN
# ─────────────────────────────────────────────
def _upload_form(key_prefix: str):
    """Affiche le formulaire d'upload et retourne (df, pays, entreprise, periode) ou None."""
    col1, col2 = st.columns([2, 1])

    with col1:
        st.markdown('<div class="upload-zone">', unsafe_allow_html=True)
        uploaded = st.file_uploader(
            "📂 Téléchargez votre balance des comptes (Excel .xlsx)",
            type=["xlsx", "xls"],
            key=f"{key_prefix}_upload",
            help="Format Sage 100 ou tout logiciel SYSCOHADA — balance 6 colonnes"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        pays = st.selectbox(
            "🌍 Pays",
            options=list(PAYS_UEMOA.keys()),
            key=f"{key_prefix}_pays"
        )
        entreprise = st.text_input(
            "🏢 Nom de l'entreprise",
            value="",
            key=f"{key_prefix}_entreprise",
            placeholder="Ex : DARLING SARL"
        )
        periode = st.text_input(
            "📅 Période",
            value="",
            key=f"{key_prefix}_periode",
            placeholder="Ex : Exercice 2024"
        )

    if uploaded is None:
        return None

    try:
        df = parse_balance(uploaded)
        if df.empty:
            st.error("❌ Aucune donnée comptable détectée. Vérifiez le format de la balance.")
            return None
        st.success(f"✅ Balance chargée — {len(df)} comptes détectés")
        return df, pays, entreprise or "Entreprise", periode or "Période non précisée"
    except Exception as e:
        st.error(f"❌ Erreur de lecture : {e}")
        return None


# ─────────────────────────────────────────────
#  PAGE 1 : TABLEAU DE BORD FISCAL
# ─────────────────────────────────────────────
def page_dashboard():
    """
    Page Streamlit — Tableau de Bord Fiscal Mensuel.
    À appeler depuis votre app principale.
    """
    _inject_css()

    # En-tête
    st.markdown("""
    <div class="smd-header">
        <h1>📊 Tableau de Bord Fiscal Mensuel</h1>
        <p>SMD Consulting — Superviseur IA Comptable & Fiscal UEMOA / OHADA</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("Uploadez votre balance des comptes pour générer instantanément le tableau de bord fiscal selon les règles fiscales du pays sélectionné.")

    # Formulaire
    result_data = _upload_form("tdb")
    if result_data is None:
        _info_pays()
        return

    df, pays, entreprise, periode = result_data

    with st.spinner("⚙️ Calcul des indicateurs fiscaux..."):
        result = compute_fiscal(df, pays, entreprise, periode)

    # ── KPIs principaux ──
    st.markdown("---")
    st.markdown(f"### 📈 Indicateurs — {entreprise} | {pays} | {periode}")

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        _kpi("Chiffre d'Affaires", _fmt(result["ca"]), "bleu")
    with c2:
        couleur = "vert" if result["resultat"] > 0 else "rouge"
        _kpi("Résultat Net", _fmt(result["resultat"]), couleur)
    with c3:
        couleur = "rouge" if result["tva_nette"] > 0 else "vert"
        _kpi(f"TVA/IGV Nette Due ({result['tva_taux']}%)", _fmt(result["tva_nette"]), couleur)
    with c4:
        couleur = "orange" if result["is_solde"] > 0 else "vert"
        _kpi(f"IS/IRPC Résiduel ({result['is_taux']}%)", _fmt(result["is_solde"]), couleur)

    # ── Exposition totale ──
    if result["exposition_totale"] > 0:
        st.markdown(f"""
        <div class="exposition-box">
            <div class="label">⚠️ EXPOSITION FISCALE TOTALE ESTIMÉE</div>
            <div class="montant">{_fmt(result["exposition_totale"])}</div>
            <small style="color:#C0392B;">hors pénalités et majorations de retard {result["admin"]}</small>
        </div>
        """, unsafe_allow_html=True)

    # ── Onglets détaillés ──
    st.markdown("---")
    tab_tva, tab_is, tab_inps, tab_alertes, tab_ratios = st.tabs([
        "💰 TVA / IGV", "🏦 IS / IRPC", "👥 INPS / CNSS", "🔔 Alertes", "📐 Ratios"
    ])

    # Tab TVA
    with tab_tva:
        col_g, col_d = st.columns([1, 1])
        with col_g:
            st.markdown(f"**Taux TVA/IGV : {result['tva_taux']} % — Administration : {result['admin']}**")
            tva_df = pd.DataFrame([
                {"Élément": "TVA collectée",    "Montant (FCFA)": int(result["tva_collectee"]),  "Nature": "Passif fiscal"},
                {"Élément": "TVA déductible",   "Montant (FCFA)": int(result["tva_deductible"]), "Nature": "Crédit"},
                {"Élément": "TVA NETTE DUE",    "Montant (FCFA)": int(result["tva_nette"]),      "Nature": "À REVERSER"},
            ])
            st.dataframe(tva_df, use_container_width=True, hide_index=True)

            niveau_tva = "🔴 RISQUE ÉLEVÉ" if result["tva_nette"] > result["ca"] * 0.03 else \
                         "🟠 À SURVEILLER" if result["tva_nette"] > 0 else "🟢 CRÉDIT TVA"
            st.info(f"**Position TVA :** {niveau_tva} — Échéance : **{result['echeance_tva']}**")

        with col_d:
            st.plotly_chart(_chart_tva(result), use_container_width=True)

        # Détail TVA déductible
        if result["tva_detail"]:
            st.markdown("**Détail TVA déductible par compte :**")
            detail_df = pd.DataFrame(result["tva_detail"])
            detail_df["montant"] = detail_df["montant"].apply(int)
            detail_df.columns = ["Compte", "Libellé", "Montant (FCFA)"]
            st.dataframe(detail_df, use_container_width=True, hide_index=True)

    # Tab IS
    with tab_is:
        col_g, col_d = st.columns([1, 1])
        with col_g:
            st.markdown(f"**Taux IS/IRPC : {result['is_taux']} %**")
            is_df = pd.DataFrame([
                {"Calcul": "Résultat comptable",   "FCFA": int(result["resultat"])},
                {"Calcul": "+ Réintégrations",     "FCFA": int(result["reintegrations"])},
                {"Calcul": "= Résultat fiscal",    "FCFA": int(result["resultat_fiscal"])},
                {"Calcul": f"IS ({result['is_taux']}%)", "FCFA": int(result["is_corrige"])},
                {"Calcul": "IS versé",             "FCFA": int(result["is_paye"])},
                {"Calcul": "IS RÉSIDUEL",          "FCFA": int(result["is_solde"])},
            ])
            st.dataframe(is_df, use_container_width=True, hide_index=True)

            if result["imf"] > 0:
                st.warning(f"⚠️ IMF (minimum forfaitaire) : {_fmt(result['imf'])} — IS à payer = max(IS, IMF)")
            st.info(f"**Échéance déclaration IS :** {result['echeance_is']}")
        with col_d:
            st.plotly_chart(_chart_is(result), use_container_width=True)

    # Tab INPS
    with tab_inps:
        col_g, col_d = st.columns([1, 1])
        with col_g:
            st.markdown(f"**Part patronale : {result['inps_patron_taux']} % — Part salariale : {result['inps_sal_taux']} %**")
            inps_df = pd.DataFrame([
                {"Élément": "Salaires bruts",           "FCFA": int(result["salaires_bruts"])},
                {"Élément": f"Patronal ({result['inps_patron_taux']}%)", "FCFA": int(result["inps_patron_du"])},
                {"Élément": f"Salarial ({result['inps_sal_taux']}%)",   "FCFA": int(result["inps_sal_du"])},
                {"Élément": "TOTAL INPS attendu",       "FCFA": int(result["inps_total_du"])},
                {"Élément": "INPS versé",               "FCFA": int(result["inps_paye"])},
                {"Élément": "Écart à vérifier",         "FCFA": int(result["inps_ecart"])},
            ])
            st.dataframe(inps_df, use_container_width=True, hide_index=True)
            st.info(f"**Échéance INPS/CNSS :** {result['echeance_inps']}")
        with col_d:
            st.plotly_chart(_chart_inps(result), use_container_width=True)

    # Tab Alertes
    with tab_alertes:
        if not result["alertes"]:
            st.success("✅ Aucune alerte fiscale majeure détectée.")
        else:
            st.markdown(f"**{len(result['alertes'])} alerte(s) identifiée(s) :**")
            for a in sorted(result["alertes"], key=lambda x: {"ÉLEVÉ": 0, "MODÉRÉ": 1}.get(x["niveau"], 2)):
                _alerte_box(a["titre"], a["montant"], a["detail"], a["action"], a["niveau"])

            # Graphique exposition
            fig_exp = _chart_exposition(result)
            if fig_exp:
                st.plotly_chart(fig_exp, use_container_width=True)

    # Tab Ratios
    with tab_ratios:
        ratios_df = pd.DataFrame([
            {"Ratio": "TVA nette / CA",       "Valeur": f"{result['ratio_tva_ca']:.2f} %",      "Alerte": "Normal si < taux TVA × marge"},
            {"Ratio": "Pression fiscale / CA","Valeur": f"{result['pression_fiscale']:.2f} %",    "Alerte": "⚠️ Élevé si > 35 %"},
            {"Ratio": "Capital / CA",         "Valeur": f"{result['ratio_capital_ca']:.2f} %",    "Alerte": "⚠️ Critique si < 1 %"},
            {"Ratio": "Résultat / CA",        "Valeur": f"{result['resultat']/result['ca']*100:.2f} %" if result['ca'] else "N/A", "Alerte": "Marge nette"},
        ])
        st.dataframe(ratios_df, use_container_width=True, hide_index=True)

        c_rat1, c_rat2, c_rat3 = st.columns(3)
        with c_rat1:
            couleur = "rouge" if result["ratio_capital_ca"] < 1 else "vert"
            _kpi("Capital / CA", f"{result['ratio_capital_ca']:.2f} %", couleur)
        with c_rat2:
            couleur = "rouge" if result["pression_fiscale"] > 35 else "vert"
            _kpi("Pression fiscale", f"{result['pression_fiscale']:.2f} %", couleur)
        with c_rat3:
            _kpi("TVA nette / CA", f"{result['ratio_tva_ca']:.2f} %", "bleu")

    # ── Téléchargements ──
    st.markdown("---")
    st.markdown("### 📥 Télécharger les Rapports")
    dl_col1, dl_col2 = st.columns(2)

    with dl_col1:
        excel_buf = generate_excel(result)
        st.download_button(
            label="⬇️ Tableau de bord Excel (.xlsx)",
            data=excel_buf,
            file_name=f"Tableau_Bord_Fiscal_{entreprise.replace(' ','_')}_{pays}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

    with dl_col2:
        word_buf = generate_word(result)
        if word_buf:
            st.download_button(
                label="⬇️ Rapport Word (.docx)",
                data=word_buf,
                file_name=f"Rapport_Risque_Fiscal_{entreprise.replace(' ','_')}_{pays}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.info("Installez `python-docx` pour activer l'export Word : `pip install python-docx`")

    # Pied de page
    st.markdown("---")
    st.caption(
        "SMD Consulting — Superviseur IA Comptable & Fiscal UEMOA/OHADA v2.0  |  "
        "Ce tableau de bord est produit sur la base de la balance fournie. "
        "Il ne constitue pas un audit légal."
    )


# ─────────────────────────────────────────────
#  PAGE 2 : ANALYSE DU RISQUE FISCAL
# ─────────────────────────────────────────────
def page_risque_fiscal():
    """
    Page Streamlit — Analyse complète du Risque Fiscal.
    À appeler depuis votre app principale.
    """
    _inject_css()

    st.markdown("""
    <div class="smd-header">
        <h1>🔍 Analyse du Risque Fiscal</h1>
        <p>SMD Consulting — Détection des risques fiscaux & préparation aux contrôles DGCI / DGI</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("Analyse approfondie des risques fiscaux selon le cadre légal du pays. Identifie les expositions, calcule les montants à régulariser et priorise les actions.")

    result_data = _upload_form("risque")
    if result_data is None:
        _info_controle()
        return

    df, pays, entreprise, periode = result_data

    with st.spinner("🔎 Analyse des risques en cours..."):
        result = compute_fiscal(df, pays, entreprise, periode)

    # ── Résumé exécutif ──
    st.markdown("---")

    # Score de risque global
    n_eleve  = sum(1 for a in result["alertes"] if a["niveau"] == "ÉLEVÉ")
    n_modere = sum(1 for a in result["alertes"] if a["niveau"] == "MODÉRÉ")
    score    = n_eleve * 3 + n_modere

    if score >= 6:
        niveau_global, couleur_global, emoji_global = "ÉLEVÉ", COULEUR_ROUGE, "🔴"
    elif score >= 3:
        niveau_global, couleur_global, emoji_global = "MODÉRÉ", COULEUR_ORANGE, "🟠"
    else:
        niveau_global, couleur_global, emoji_global = "FAIBLE", COULEUR_VERT, "🟢"

    col_score, col_details = st.columns([1, 2])
    with col_score:
        st.markdown(f"""
        <div style="background:{couleur_global}; color:white; border-radius:12px;
                    padding:1.5rem; text-align:center; height:140px;
                    display:flex; flex-direction:column; justify-content:center;">
            <div style="font-size:3rem;">{emoji_global}</div>
            <div style="font-size:1.2rem; font-weight:700;">RISQUE {niveau_global}</div>
            <div style="opacity:0.85; font-size:0.85rem;">{n_eleve} élevé · {n_modere} modéré</div>
        </div>
        """, unsafe_allow_html=True)

    with col_details:
        c1, c2 = st.columns(2)
        with c1:
            _kpi("Exposition totale", _fmt(result["exposition_totale"]), "rouge")
            _kpi("TVA/IGV nette due", _fmt(result["tva_nette"]), "rouge" if result["tva_nette"] > 0 else "vert")
        with c2:
            _kpi("IS/IRPC résiduel", _fmt(result["is_solde"]), "orange" if result["is_solde"] > 0 else "vert")
            _kpi("INPS écart", _fmt(result["inps_ecart"]), "orange" if result["inps_ecart"] > 0 else "vert")

    # ── Tableau des risques ──
    st.markdown("---")
    st.markdown("### 🗂️ Registre des Risques Fiscaux")

    if result["alertes"]:
        rows = []
        for a in sorted(result["alertes"], key=lambda x: {"ÉLEVÉ": 0, "MODÉRÉ": 1}.get(x["niveau"], 2)):
            rows.append({
                "Risque":   a["titre"],
                "Montant":  _fmt(a["montant"]),
                "Niveau":   f"{NIVEAU_EMOJI.get(a['niveau'], '')} {a['niveau']}",
                "Action":   a["action"][:80] + "…" if len(a["action"]) > 80 else a["action"],
            })
        df_alertes = pd.DataFrame(rows)

        # Coloration conditionnelle
        def highlight_niveau(row):
            colors = []
            for col in row.index:
                if "ÉLEVÉ" in str(row["Niveau"]):
                    colors.append("background-color: #FADBD8")
                elif "MODÉRÉ" in str(row["Niveau"]):
                    colors.append("background-color: #FDEBD0")
                else:
                    colors.append("")
            return colors

        styled = df_alertes.style.apply(highlight_niveau, axis=1)
        st.dataframe(styled, use_container_width=True, hide_index=True)

        # Graphique
        fig_exp = _chart_exposition(result)
        if fig_exp:
            st.plotly_chart(fig_exp, use_container_width=True)
    else:
        st.success("✅ Aucun risque fiscal majeur détecté sur la base de la balance fournie.")

    # ── Détail par impôt ──
    st.markdown("---")
    st.markdown("### 📋 Analyse Détaillée par Impôt")

    with st.expander(f"💰 TVA / IGV — Taux {result['tva_taux']} % ({result['admin']})", expanded=True):
        c1, c2, c3 = st.columns(3)
        c1.metric("Collectée", _fmt(result["tva_collectee"]))
        c2.metric("Déductible", _fmt(result["tva_deductible"]))
        c3.metric("Nette due", _fmt(result["tva_nette"]),
                  delta=f"{'⚠️ À régulariser' if result['tva_nette'] > 0 else '✅ Crédit'}")
        st.info(f"📅 Échéance déclaration mensuelle : **{result['echeance_tva']}**")
        st.markdown(
            "**Recommandation :** Rapprocher les soldes TVA comptables avec les déclarations "
            f"mensuelles déposées auprès de la {result['admin']}. "
            "Tout écart non justifié est passible d'amendes et majorations."
        )

    with st.expander(f"🏦 IS / IRPC — Taux {result['is_taux']} %", expanded=True):
        c1, c2, c3 = st.columns(3)
        c1.metric("Résultat fiscal", _fmt(result["resultat_fiscal"]))
        c2.metric("IS calculé", _fmt(result["is_corrige"]))
        c3.metric("IS résiduel", _fmt(result["is_solde"]),
                  delta=f"{'⚠️ À verser' if result['is_solde'] > 0 else '✅ OK'}")
        if result["reintegrations"] > 0:
            st.warning(f"⚠️ Réintégrations fiscales à opérer : {_fmt(result['reintegrations'])} (dons, pénalités non déductibles)")
        st.info(f"📅 Échéance déclaration IS : **{result['echeance_is']}**")

    with st.expander(f"👥 INPS / CNSS — Patronal {result['inps_patron_taux']} % | Salarial {result['inps_sal_taux']} %"):
        c1, c2, c3 = st.columns(3)
        c1.metric("Salaires bruts", _fmt(result["salaires_bruts"]))
        c2.metric("INPS total attendu", _fmt(result["inps_total_du"]))
        c3.metric("Écart à vérifier", _fmt(result["inps_ecart"]))
        st.info(f"📅 Échéance cotisations : **{result['echeance_inps']}**")
        if result["avances_personnel"] > 0:
            st.warning(f"⚠️ Avances au personnel non remboursées : {_fmt(result['avances_personnel'])} — risque requalification avantage en nature")

    # ── Plan d'action ──
    st.markdown("---")
    st.markdown("### 📌 Plan d'Action Prioritaire")

    actions_urgentes   = [a for a in result["alertes"] if a["niveau"] == "ÉLEVÉ"]
    actions_importantes = [a for a in result["alertes"] if a["niveau"] == "MODÉRÉ"]

    if actions_urgentes:
        st.error(f"**🔴 ACTIONS URGENTES ({len(actions_urgentes)}) :**")
        for i, a in enumerate(actions_urgentes, 1):
            st.markdown(f"**{i}.** {a['action']}")

    if actions_importantes:
        st.warning(f"**🟠 ACTIONS IMPORTANTES ({len(actions_importantes)}) :**")
        for i, a in enumerate(actions_importantes, 1):
            st.markdown(f"**{i}.** {a['action']}")

    st.info(
        f"**🔵 PRÉVENTIF :** Prendre contact avec un avocat fiscaliste local en {pays} "
        f"pour initier une régularisation spontanée auprès de la {result['admin']} "
        "avant tout contrôle fiscal."
    )

    # ── Téléchargements ──
    st.markdown("---")
    st.markdown("### 📥 Télécharger les Rapports")
    dl1, dl2 = st.columns(2)

    with dl1:
        excel_buf = generate_excel(result)
        st.download_button(
            label="⬇️ Tableau de bord Excel (.xlsx)",
            data=excel_buf,
            file_name=f"Risque_Fiscal_{entreprise.replace(' ','_')}_{pays}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

    with dl2:
        word_buf = generate_word(result)
        if word_buf:
            st.download_button(
                label="⬇️ Rapport Word (.docx)",
                data=word_buf,
                file_name=f"Analyse_Risque_Fiscal_{entreprise.replace(' ','_')}_{pays}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.info("Installez `python-docx` pour activer l'export Word.")

    st.caption(
        "SMD Consulting — Ce rapport est produit sur la base de la balance fournie. "
        "Il ne constitue pas un audit légal. "
        f"Administration fiscale compétente : {result['admin']} — {pays}."
    )


# ─────────────────────────────────────────────
#  BLOCS D'INFO (affiché avant upload)
# ─────────────────────────────────────────────
def _info_pays():
    st.markdown("---")
    st.markdown("### 🌍 Pays UEMOA couverts")
    cols = st.columns(4)
    pays_list = list(PAYS_UEMOA.items())
    for i, (pays, params) in enumerate(pays_list):
        with cols[i % 4]:
            st.markdown(f"""
            <div style="background:#EBF5FB; border-radius:8px; padding:0.7rem;
                        margin-bottom:0.5rem; border-left:4px solid #1A5276;">
                <strong>{pays}</strong><br>
                <small>TVA {params['tva_taux']}% · IS {params['is_taux']}%<br>
                {params['admin']}</small>
            </div>
            """, unsafe_allow_html=True)


def _info_controle():
    st.markdown("---")
    st.markdown("### ℹ️ Ce module analyse :")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("""
        **💰 Position TVA/IGV**
        - TVA collectée vs déductible
        - IGV nette potentiellement due
        - Détail par compte de TVA
        """)
    with c2:
        st.markdown("""
        **🏦 IS / IRPC**
        - Calcul du résultat fiscal
        - Réintégrations (dons, pénalités)
        - Acomptes versés vs IS dû
        """)
    with c3:
        st.markdown("""
        **👥 INPS / CNSS**
        - Cotisations patronale et salariale
        - Vérification des versements
        - Avances au personnel
        """)


# ─────────────────────────────────────────────
#  PAGE : ANALYSE FACTURE SYSCOHADA
# ─────────────────────────────────────────────

# Référentiel d'imputation SYSCOHADA par nature de charge/produit
_CHARGES_SYSCOHADA = {
    "Marchandises (achat pour revente)":         ("601", "Achats de marchandises"),
    "Matières premières":                         ("602", "Achats de matières premières"),
    "Fournitures de bureau / consommables":       ("6052", "Fournitures de bureau"),
    "Carburant / lubrifiant":                     ("6053", "Fournitures d'entretien et de petit équipement"),
    "Sous-traitance générale":                    ("604", "Achats d'études et prestations de services"),
    "Loyer / location immobilier":                ("6221", "Loyers et charges locatives"),
    "Électricité / eau / gaz":                   ("6053", "Eau, électricité, gaz"),
    "Téléphone / internet":                       ("6252", "Frais de téléphone et internet"),
    "Transport / fret":                           ("624", "Transports de biens et transports collectifs"),
    "Honoraires (expert-comptable, avocat...)":   ("6221", "Honoraires"),
    "Publicité / communication":                  ("627", "Publicité, publications, relations publiques"),
    "Assurances":                                 ("625", "Primes d'assurances"),
    "Frais bancaires":                            ("671", "Frais bancaires"),
    "Matériel informatique (immobilisation)":     ("244", "Matériel informatique"),
    "Mobilier / matériel de bureau (immo.)":      ("244", "Matériel et mobilier de bureau"),
    "Véhicule (immobilisation)":                  ("245", "Matériel de transport"),
    "Travaux / aménagements (immo.)":             ("232", "Agencements et installations"),
    "Salaires et traitements":                    ("661", "Rémunérations directes versées"),
    "Autre charge (préciser en libellé)":         ("658", "Autres charges diverses"),
}

_PRODUITS_SYSCOHADA = {
    "Ventes de marchandises":                     ("701", "Ventes de marchandises"),
    "Ventes de produits finis":                   ("702", "Ventes de produits finis"),
    "Prestations de services":                    ("706", "Services vendus"),
    "Travaux facturés":                           ("703", "Travaux facturés"),
    "Loyers perçus":                              ("771", "Revenus des immeubles"),
    "Produits financiers (intérêts)":             ("772", "Revenus des valeurs mobilières"),
    "Autre produit (préciser en libellé)":        ("758", "Autres produits divers"),
}

_COMPTES_TVA_COL  = "4431"   # TVA collectée (ventes)
_COMPTES_TVA_DED  = "4454"   # TVA déductible (achats)
_COMPTE_FOUR      = "401"    # Fournisseurs
_COMPTE_CLIENT    = "411"    # Clients


def _generer_excel_facture(lignes: list, meta: dict) -> BytesIO:
    """Génère un fichier Excel avec le journal de la facture."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    wb = Workbook()
    ws = wb.active
    ws.title = "Journal Facture"

    bleu  = "1A5276"
    blanc = "FFFFFF"
    gris  = "F2F2F2"

    # En-tête document
    ws.merge_cells("A1:G1")
    ws["A1"] = f"ANALYSE FACTURE — {meta.get('entreprise','').upper()}"
    ws["A1"].font = Font(bold=True, color=blanc, size=13)
    ws["A1"].fill = PatternFill("solid", fgColor=bleu)
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A2:G2")
    ws["A2"] = f"Facture N° {meta.get('num_facture','')} — {meta.get('date_facture','')} — {meta.get('pays','')}"
    ws["A2"].font = Font(italic=True)
    ws["A2"].alignment = Alignment(horizontal="center")

    # En-têtes colonnes
    headers = ["Date", "Journal", "N° Compte", "Libellé", "Débit (FCFA)", "Crédit (FCFA)", "Ref. Facture"]
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=4, column=col, value=h)
        c.font = Font(bold=True, color=blanc)
        c.fill = PatternFill("solid", fgColor=bleu)
        c.alignment = Alignment(horizontal="center")

    # Lignes d'écriture
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"),  bottom=Side(style="thin")
    )
    for i, ligne in enumerate(lignes):
        row = 5 + i
        fill = PatternFill("solid", fgColor=gris) if i % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
        vals = [ligne.get("date"), ligne.get("journal"), ligne.get("compte"),
                ligne.get("libelle"), ligne.get("debit"), ligne.get("credit"),
                ligne.get("ref")]
        for col, v in enumerate(vals, 1):
            c = ws.cell(row=row, column=col, value=v)
            c.fill = fill
            c.border = thin
            if col in (5, 6) and isinstance(v, (int, float)):
                c.number_format = "#,##0"

    # Totaux
    row_total = 5 + len(lignes)
    ws.cell(row=row_total, column=4, value="TOTAL").font = Font(bold=True)
    total_deb = sum(l.get("debit", 0) or 0 for l in lignes)
    total_cre = sum(l.get("credit", 0) or 0 for l in lignes)
    ws.cell(row=row_total, column=5, value=total_deb).font = Font(bold=True)
    ws.cell(row=row_total, column=6, value=total_cre).font = Font(bold=True)

    # Largeurs
    for col, w in zip("ABCDEFG", [12, 10, 14, 40, 18, 18, 16]):
        ws.column_dimensions[col].width = w

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _generer_word_facture(lignes: list, meta: dict) -> BytesIO:
    """Génère un rapport Word de l'analyse facture."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = Document()
    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Titre
    titre = doc.add_heading("ANALYSE FACTURE — SYSCOHADA", level=1)
    titre.runs[0].font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    # Méta
    doc.add_paragraph(f"Entreprise   : {meta.get('entreprise','')}")
    doc.add_paragraph(f"Facture N°   : {meta.get('num_facture','')}  |  Date : {meta.get('date_facture','')}")
    doc.add_paragraph(f"Fournisseur/Client : {meta.get('tiers','')}  |  Pays : {meta.get('pays','')}")
    doc.add_paragraph(f"Type         : {meta.get('type_facture','')}  |  Nature : {meta.get('nature','')}")
    doc.add_paragraph(f"Montant HT   : {_fmt(meta.get('ht',0))}  |  TVA ({meta.get('taux_tva',0)}%) : {_fmt(meta.get('tva',0))}  |  TTC : {_fmt(meta.get('ttc',0))}")
    doc.add_paragraph("")

    doc.add_heading("ÉCRITURE COMPTABLE (Journal SYSCOHADA)", level=2)

    # Tableau
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["N° Compte", "Libellé", "Débit (FCFA)", "Crédit (FCFA)", "Ref."]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True

    for ligne in lignes:
        row = table.add_row().cells
        row[0].text = str(ligne.get("compte", ""))
        row[1].text = str(ligne.get("libelle", ""))
        row[2].text = f"{ligne.get('debit', '') or '':,.0f}".replace(",", " ") if ligne.get("debit") else ""
        row[3].text = f"{ligne.get('credit', '') or '':,.0f}".replace(",", " ") if ligne.get("credit") else ""
        row[4].text = str(ligne.get("ref", ""))

    doc.add_paragraph("")
    doc.add_heading("RAPPEL DES RÈGLES D'IMPUTATION SYSCOHADA", level=2)
    if meta.get("type_facture") == "Facture d'achat (fournisseur)":
        doc.add_paragraph("• Classe 6 (charges) ou Classe 2 (immobilisations) au DÉBIT")
        doc.add_paragraph(f"• Compte {_COMPTES_TVA_DED} — TVA déductible au DÉBIT (si assujetti TVA)")
        doc.add_paragraph(f"• Compte {_COMPTE_FOUR} — Fournisseurs au CRÉDIT (dette fournisseur)")
    else:
        doc.add_paragraph("• Classe 4 — Clients (411) au DÉBIT (créance client)")
        doc.add_paragraph("• Classe 7 (produits) au CRÉDIT")
        doc.add_paragraph(f"• Compte {_COMPTES_TVA_COL} — TVA collectée au CRÉDIT (si assujetti TVA)")

    doc.add_paragraph("")
    doc.add_paragraph("SMD Consulting — Superviseur IA Comptable & Fiscal UEMOA/OHADA", style="Caption")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def page_analyse_facture():
    """Page Streamlit : Analyse et imputation comptable d'une facture SYSCOHADA."""
    _inject_css()

    st.title("🧾 Analyse Facture — Imputation SYSCOHADA")
    st.markdown("Saisissez les informations de votre facture. L'assistant génère l'écriture comptable conforme au Plan Comptable SYSCOHADA.")
    st.divider()

    # ── Formulaire de saisie ──────────────────────────────────────
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📋 Informations générales")
        entreprise   = st.text_input("Nom de l'entreprise", placeholder="DARLING SARL")
        tiers        = st.text_input("Fournisseur / Client", placeholder="SOCIÉTÉ ABC")
        num_facture  = st.text_input("N° de facture", placeholder="F-2025-001")
        date_facture = st.date_input("Date de la facture")
        pays         = st.selectbox("Pays UEMOA", list(PAYS_UEMOA.keys()))

    with col2:
        st.subheader("💰 Détails financiers")
        type_facture = st.selectbox("Type de facture", [
            "Facture d'achat (fournisseur)",
            "Facture de vente (client)",
        ])
        if type_facture == "Facture d'achat (fournisseur)":
            catalogue = _CHARGES_SYSCOHADA
            label_nature = "Nature de la charge / immobilisation"
        else:
            catalogue = _PRODUITS_SYSCOHADA
            label_nature = "Nature du produit / service vendu"

        nature   = st.selectbox(label_nature, list(catalogue.keys()))
        libelle_libre = st.text_input("Libellé libre (optionnel)", placeholder="Ex: Loyer bureau Dakar nov. 2025")
        montant_ht    = st.number_input("Montant HT (FCFA)", min_value=0.0, step=1000.0, format="%.0f")
        taux_tva_pays = PAYS_UEMOA[pays]["tva_taux"]
        assujetti_tva = st.checkbox(f"Assujetti à la TVA ({taux_tva_pays:.0f}% — {pays})", value=True)
        if assujetti_tva:
            taux_tva = st.number_input("Taux TVA (%)", value=taux_tva_pays, min_value=0.0, max_value=30.0, step=0.5)
        else:
            taux_tva = 0.0

    # ── Calculs ──────────────────────────────────────────────────
    montant_tva = round(montant_ht * taux_tva / 100, 0)
    montant_ttc = montant_ht + montant_tva

    st.divider()
    c1, c2, c3 = st.columns(3)
    c1.metric("Montant HT",  _fmt(montant_ht))
    c2.metric(f"TVA ({taux_tva:.0f}%)", _fmt(montant_tva))
    c3.metric("Montant TTC", _fmt(montant_ttc))

    # ── Génération de l'écriture ─────────────────────────────────
    if st.button("📒 Générer l'écriture comptable", type="primary", use_container_width=True):
        if montant_ht <= 0:
            st.warning("⚠️ Veuillez saisir un montant HT supérieur à 0.")
            st.stop()

        compte_nature, libelle_nature = catalogue[nature]
        libelle_ecriture = libelle_libre if libelle_libre else libelle_nature
        ref = num_facture or "FA-001"
        date_str = date_facture.strftime("%d/%m/%Y")

        # Construction des lignes du journal
        lignes = []

        if type_facture == "Facture d'achat (fournisseur)":
            # DÉBIT : compte charge/immo
            lignes.append({"date": date_str, "journal": "ACH", "compte": compte_nature,
                            "libelle": libelle_ecriture, "debit": montant_ht, "credit": None, "ref": ref})
            # DÉBIT : TVA déductible (si assujetti)
            if assujetti_tva and montant_tva > 0:
                lignes.append({"date": date_str, "journal": "ACH", "compte": _COMPTES_TVA_DED,
                                "libelle": f"TVA déductible — {libelle_ecriture}",
                                "debit": montant_tva, "credit": None, "ref": ref})
            # CRÉDIT : fournisseur
            lignes.append({"date": date_str, "journal": "ACH", "compte": _COMPTE_FOUR,
                            "libelle": f"Fournisseur — {tiers}", "debit": None,
                            "credit": montant_ttc, "ref": ref})
        else:
            # DÉBIT : client
            lignes.append({"date": date_str, "journal": "VTE", "compte": _COMPTE_CLIENT,
                            "libelle": f"Client — {tiers}", "debit": montant_ttc,
                            "credit": None, "ref": ref})
            # CRÉDIT : produit
            lignes.append({"date": date_str, "journal": "VTE", "compte": compte_nature,
                            "libelle": libelle_ecriture, "debit": None,
                            "credit": montant_ht, "ref": ref})
            # CRÉDIT : TVA collectée (si assujetti)
            if assujetti_tva and montant_tva > 0:
                lignes.append({"date": date_str, "journal": "VTE", "compte": _COMPTES_TVA_COL,
                                "libelle": f"TVA collectée — {libelle_ecriture}",
                                "debit": None, "credit": montant_tva, "ref": ref})

        # Vérification équilibre
        total_deb = sum(l["debit"]  or 0 for l in lignes)
        total_cre = sum(l["credit"] or 0 for l in lignes)
        equilibre = abs(total_deb - total_cre) < 1

        # ── Affichage tableau ─────────────────────────────────────
        st.subheader("📒 Écriture comptable SYSCOHADA")
        df_journal = pd.DataFrame(lignes)[["date","journal","compte","libelle","debit","credit","ref"]]
        df_journal.columns = ["Date", "Journal", "N° Compte", "Libellé", "Débit", "Crédit", "Réf."]
        st.dataframe(
            df_journal.style.format({
                "Débit":  lambda x: f"{x:,.0f}".replace(",", " ") if pd.notna(x) else "",
                "Crédit": lambda x: f"{x:,.0f}".replace(",", " ") if pd.notna(x) else "",
            }),
            use_container_width=True, hide_index=True
        )

        # Équilibre
        if equilibre:
            st.success(f"✅ Écriture équilibrée — Débit = Crédit = **{_fmt(total_deb)}**")
        else:
            st.error(f"❌ Déséquilibre détecté : Débit {_fmt(total_deb)} ≠ Crédit {_fmt(total_cre)}")

        # ── Rappel des comptes ────────────────────────────────────
        with st.expander("📌 Détail des comptes SYSCOHADA utilisés"):
            est_achat    = (type_facture == "Facture d'achat (fournisseur)")
            role_nature  = "Charge / Immobilisation" if est_achat else "Produit"
            cpt_tva      = _COMPTES_TVA_DED if est_achat else _COMPTES_TVA_COL
            lib_tva      = "TVA déductible sur achats" if est_achat else "TVA collectée sur ventes"
            cpt_tiers    = _COMPTE_FOUR if est_achat else _COMPTE_CLIENT
            lib_tiers    = "Fournisseurs" if est_achat else "Clients"
            st.markdown(f"""
| N° Compte | Libellé | Rôle |
|-----------|---------|------|
| **{compte_nature}** | {libelle_nature} | {role_nature} |
| **{cpt_tva}** | {lib_tva} | TVA |
| **{cpt_tiers}** | {lib_tiers} | Tiers |
""")

        # ── Exports ───────────────────────────────────────────────
        st.divider()
        st.subheader("📥 Télécharger l'écriture")
        meta = {
            "entreprise": entreprise, "tiers": tiers, "num_facture": num_facture,
            "date_facture": date_str, "pays": pays, "type_facture": type_facture,
            "nature": nature, "ht": montant_ht, "tva": montant_tva,
            "ttc": montant_ttc, "taux_tva": taux_tva,
        }
        col_xl, col_wd = st.columns(2)
        with col_xl:
            buf_xl = _generer_excel_facture(lignes, meta)
            st.download_button(
                "📊 Télécharger Excel",
                data=buf_xl,
                file_name=f"Facture_{num_facture or 'analyse'}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
        with col_wd:
            buf_wd = _generer_word_facture(lignes, meta)
            st.download_button(
                "📄 Télécharger Word",
                data=buf_wd,
                file_name=f"Facture_{num_facture or 'analyse'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )


# ─────────────────────────────────────────────
#  APP STANDALONE (test rapide)
# ─────────────────────────────────────────────
if __name__ == "__main__":
    st.set_page_config(
        page_title="SMD Consulting — Fiscal UEMOA",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    with st.sidebar:
        st.image("https://via.placeholder.com/200x60/1A5276/FFFFFF?text=SMD+Consulting",
                 use_column_width=True)
        st.markdown("---")
        module = st.radio(
            "🗂️ Module",
            ["📊 Tableau de Bord Fiscal", "🔍 Analyse du Risque Fiscal", "🧾 Analyse Facture"],
            label_visibility="collapsed"
        )
        st.markdown("---")
        st.markdown("**SMD Consulting**")
        st.markdown("Superviseur IA Comptable & Fiscal")
        st.markdown("UEMOA / OHADA — v2.0")

    if module == "📊 Tableau de Bord Fiscal":
        page_dashboard()
    elif module == "🔍 Analyse du Risque Fiscal":
        page_risque_fiscal()
    else:
        page_analyse_facture()
        page_risque_fiscal()
    else:
        page_analyse_facture()
